# 文件路径处理
import sys
from pathlib import Path

# excel处理
import polars as pl

# word处理
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# PDF处理
from docx2pdf import convert

"""
    这个文件的作用是: 将excel中的数据进行处理, 
    然后生成对应的 word 文件, 最后转成 PDF 文件
    实现逻辑: 每一个处理都是对单个文件的处理, 然后在套上一个方法: 批处理调度器
    本质逻辑还是单个文件处理. 主要是处理一些简单的excel表格, 
    若是特殊表格, 那么只能是特殊情况特殊处理了.
    excel --> word 方法都将其单独整理出来了.
    word --> PDF 方法都将其单独整理出来了.
    所以也可以直接调用, 该有的说明都有.
    
    # TODO
        还差word边界情况的处理没有完成
"""


"""
TODO:
    我在想这个函数需要传递的参数是不是太多了,
    所以我后续会进行简化.
"""


# 这里的两个代码是进行将 excel 进行批量整理的代码
def excel_data_process(
    input_file: str | Path,  # 支持字符串或Path
    output_file: str | Path,  # 支持字符串或Path
    header_row: int,
    judge_column: str,
    keep_row: list[str],
    keep_column: list[str],
) -> bool:
    """
    关于这个方法:  是用来整理学校发的excel (必须是.xlsx后缀) 表中的数据, 并进行清洗,
    是一个非常简单的任务处理, 并不是非常专业的处理方式, 但是我想用来处理学校的excel绝对是足够了.
    可以根据需求决定是否需要转换成: word, 或者进一步转换成: PDF,
    后面会有对应的提示.
    :param input_file: 必须是一个文件的路径(包括名字)
    :param output_file: 必须是保存文件的路径(包括名字)
    :param header_row: 第几行作为各个列的名字 (从 1 开始计数)
    :param judge_column: 根据这一列来判断行中的元素是不是需要保留.
    :param keep_row: 存放你想要保留的行
    :param keep_column: 需要保留的列.
    :return: bool 类型的值
    """

    try:
        df = pl.read_excel(input_file, read_options={"header_row": header_row - 1})
        ans = df.filter(pl.col(judge_column).is_in(keep_row)).select(
            pl.col(keep_column)
        )
        ans.write_excel(output_file)
        return True
    except Exception as e:
        print(f"文件操作出现错误: {e}")
        return False


def batch_excel_data_process(
    input_folder: str,
    judge_column: str,
    keep_row: list[str],
    keep_column: list[str],
    header_row: int,
    output_folder: str,
) -> None:
    """
    关于这个方法:  是用来批量整理学校发的excel (必须是.xlsx后缀) 表中的数据, 并进行清洗,
    是一个非常简单的任务处理, 并不是非常专业的处理方式, 但是我想用来处理学校的excel绝对是足够了.
    可以根据需求决定是否需要转换成: word, 或者进一步转换成: PDF.
    后面会有对应的提示.
    :param input_folder: 必须是一个文件夹的路径
    :param judge_column: 根据这一列来判断行中的元素是不是需要保留.
    :param header_row: 第几行作为各个列的名字 (从 1 开始计数)
    :param keep_row: 存放你想要保留的行
    :param keep_column: 需要保留的列.
    :param output_folder: 必须是保存文件夹的路径
    :return: None
    """
    # 转换为 Path 对象
    input_path = Path(input_folder)

    # 设置输出路径
    output_path = Path(output_folder) if output_folder else input_path / "excel_output"
    output_path.mkdir(parents=True, exist_ok=True)

    # 获取 Excel 文件列表（Path 对象列表）
    excel_files = list(input_path.glob("*.xlsx"))

    if not excel_files:
        print(f"在文件夹 '{input_path}' 中没有找到Excel文件")
        return

    print(f"找到 {len(excel_files)} 个Excel文件")
    print(f"输出文件夹: {output_path}")
    print("-" * 50)

    for excel_file in excel_files:
        print(f"\n正在处理: {excel_file.name}")

        input_file = excel_file
        output_file = output_path / excel_file.name

        ans = excel_data_process(
            input_file,
            output_file,
            header_row,
            judge_column,
            keep_row,
            keep_column,
        )

        print(f"{excel_file.name}处理完成" if ans else f"{excel_file}处理失败")

    print("-" * 50)
    # 显示输出文件夹内容
    excel_files_finish = list(output_path.glob("*.xlsx"))
    if excel_files_finish:
        print(f"生成Excel文件数量: {len(excel_files_finish)} 个")
    else:
        print("警告: 输出文件夹中没有找到Excel文件")


# --- 这里的代码是将 excel 中的数据批量导入 word 并整理好格式的代码 ---


def set_global_font(doc, font_name):
    """设置文档全局中西文字体"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    # 获取或创建底层 XML 节点以支持中文字体
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def save_as_pretty_word(df, output_path, title_text="复习题库"):
    """将单个 DataFrame 转换为格式美观的 Word"""
    doc = Document()
    set_global_font(doc, "霞鹜文楷")

    # 1. 写入大标题
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.name = "霞鹜文楷"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "霞鹜文楷")

    records = df.to_dicts()

    for i, row in enumerate(records, 1):
        # 1. 题干 (加粗, 12pt)
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {row['题干']}")
        run.bold = True
        run.font.size = Pt(12)

        # 2. 选项 (A,B,C,D)
        for opt in ["A", "B", "C", "D"]:
            col_name = f"选项{opt}"
            # 检查列是否存在且不为空
            if col_name in row and row[col_name]:
                opt_p = doc.add_paragraph(style="List Bullet")
                opt_run = opt_p.add_run(f"{opt}. {row[col_name]}")
                opt_run.font.bold = True

        # 3. 正确答案 (深蓝色, 10pt)
        ans_p = doc.add_paragraph()
        ans_run = ans_p.add_run(f"【正确答案】：{row['正确答案']}")
        ans_run.font.color.rgb = RGBColor(0, 102, 204)
        ans_run.font.size = Pt(10)
        ans_run.font.bold = True

        # 4. 分割线
        doc.add_paragraph("-" * 80)

    doc.save(output_path)


# --- 批量处理逻辑 ---
def batch_process_folder(source_dir, output_dir):
    """
    遍历 source_dir 下所有 Excel，转换并保存到 output_dir
    """
    src_path = Path(source_dir)
    out_path = Path(output_dir)

    # 创建输出文件夹
    out_path.mkdir(parents=True, exist_ok=True)

    # 筛选所有 .xlsx 文件
    files = list(src_path.glob("*.xlsx"))

    if not files:
        print(f"❌ 错误: 在路径 {source_dir} 下没找到 .xlsx 文件")
        return

    print(f"🚀 开始转换任务，共 {len(files)} 个文件...")

    for file in files:
        try:
            # 1. 读取 Excel
            df = pl.read_excel(file)

            # 2. 确定输出路径和文档标题
            file_stem = file.stem
            target_word = out_path / f"{file_stem}.docx"

            # 3. 执行转换
            save_as_pretty_word(df, target_word, title_text=file_stem)
            print(f"✅ 已完成: {file_stem}.docx")

        except Exception as e:
            print(f"⚠️ 处理文件 {file.name} 时发生异常: {e}")


# 将一个文件夹中的 Excel 批量转为 PDF .
# 这个批量处理的方法已经完全足够了, 并不需要和上面一样进行套壳.
# docx2pdf 的 convert 已经可以处理的非常好了.
def convert_word_to_pdf(source_dir: Path, output_dir: Path):
    # 确保输出目录存在
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"🚀 开始批量转换任务: {source_dir} -> {output_dir}")

    try:
        # docx2pdf 的强大之处：可以直接传入两个目录
        # 它会自动匹配源目录下的所有 docx 并生成到目标目录
        convert(source_dir, output_dir)
        print(f"\n✨ 批量处理完成！PDF 已存入: {output_dir}")
    except Exception as e:
        print(f"⚠️ 批量转换过程中出现问题: {e}")


# 假设你的上述函数都在同一个模块中


def main():
    """
    主函数 - 提供交互式菜单选择处理流程
    """
    print("=" * 60)
    print("Excel 数据清洗与文档转换工具")
    print("=" * 60)

    while True:
        print("\n请选择要执行的操作：")
        print("1. 批量清洗 Excel 数据")
        print("3. 批量将 Word 文档转换为 PDF")
        print("4. 完整流程（Excel → Word → PDF）")
        print("5. 退出程序")

        choice = input("\n请输入选项 (1-5): ").strip()

        if choice == "1":
            process_excel_cleaning()
        elif choice == "3":
            process_word_to_pdf()
        elif choice == "4":
            process_full_workflow()
        elif choice == "5":
            print("感谢使用，再见！")
            sys.exit(0)
        else:
            print("❌ 无效选项，请重新选择")


def get_input_path(prompt: str, default: str = "") -> Path:
    """获取输入路径"""
    path_str = input(prompt + (f" [默认: {default}]" if default else ""))
    path_str = path_str.strip() if path_str.strip() else default
    return Path(path_str)


def process_excel_cleaning():
    """处理流程1：批量清洗Excel数据"""
    print("\n" + "=" * 40)
    print("Excel 数据清洗")
    print("=" * 40)

    input_folder = get_input_path("请输入Excel文件夹路径: ")
    output_folder = get_input_path(
        "请输入输出文件夹路径（回车则使用默认）: ", str(input_folder / "cleaned_excel")
    )

    # 获取清洗参数
    try:
        header_row = int(input("请输入标题行号（从1开始）: "))
        judge_column = input("请输入判断列名: ")
        keep_row_input = input("请输入要保留的行内容（用空格分隔）: ")
        keep_column_input = input("请输入要保留的列名（用空格分隔）: ")

        keep_row = keep_row_input.strip().split()
        keep_column = keep_column_input.strip().split()

        print("\n正在处理...")
        batch_excel_data_process(
            input_folder=str(input_folder),
            output_folder=str(output_folder),
            header_row=header_row,
            judge_column=judge_column,
            keep_row=keep_row,
            keep_column=keep_column,
        )

    except ValueError as e:
        print(f"❌ 输入参数错误: {e}")
    except Exception as e:
        print(f"❌ 处理过程中出现错误: {e}")


def process_word_to_pdf():
    """处理流程3：Word转PDF"""
    print("\n" + "=" * 40)
    print("Word 文档转 PDF")
    print("=" * 40)

    input_folder = get_input_path("请输入Word文件夹路径: ")
    output_folder = get_input_path(
        "请输入PDF输出文件夹路径（回车则使用默认）: ", str(input_folder / "pdf_output")
    )

    print("\n正在处理...")
    try:
        convert_word_to_pdf(
            source_dir=Path(input_folder), output_dir=Path(output_folder)
        )
    except Exception as e:
        print(f"❌ 处理过程中出现错误: {e}")


def process_full_workflow():
    """处理流程4：完整工作流（Excel → Word → PDF）"""
    print("\n" + "=" * 40)
    print("完整工作流：Excel → Word → PDF")
    print("=" * 40)

    input_folder = get_input_path("请输入Excel源文件夹路径: ")

    # 询问是否需要Excel清洗
    need_cleaning = input("是否需要先清洗Excel数据？(y/n): ").strip().lower()

    if need_cleaning == "y":
        # 清洗Excel
        cleaned_folder = input_folder / "cleaned_excel"
        print("\n第一步：Excel数据清洗")

        try:
            header_row = int(input("请输入标题行号（从1开始）: "))
            judge_column = input("请输入判断列名: ")
            keep_row_input = input("请输入要保留的行内容（用空格分隔）: ")
            keep_column_input = input("请输入要保留的列名（用空格分隔）: ")

            keep_row = keep_row_input.strip().split()
            keep_column = keep_column_input.strip().split()

            batch_excel_data_process(
                input_folder=str(input_folder),
                output_folder=str(cleaned_folder),
                header_row=header_row,
                judge_column=judge_column,
                keep_row=keep_row,
                keep_column=keep_column,
            )
            # 更新输入文件夹为清洗后的文件夹
            input_folder = cleaned_folder

        except Exception as e:
            print(f"❌ Excel清洗失败: {e}")
            return

    # Excel转Word
    word_folder = input_folder / "word_output"
    print("\n第二步：Excel转Word")
    try:
        batch_process_folder(source_dir=str(input_folder), output_dir=str(word_folder))
        print(f"✅ Word文档已保存到: {word_folder}")
    except Exception as e:
        print(f"❌ Excel转Word失败: {e}")
        return

    # Word转PDF
    pdf_folder = input_folder / "pdf_output"
    print("\n第三步：Word转PDF")
    try:
        convert_word_to_pdf(source_dir=Path(word_folder), output_dir=Path(pdf_folder))
        print(f"✅ PDF文档已保存到: {pdf_folder}")
    except Exception as e:
        print(f"❌ Word转PDF失败: {e}")
        return

    print("\n✨ 完整工作流处理完成！")
    print(f"原始文件: {input_folder}")
    print(f"Word文档: {word_folder}")
    print(f"PDF文档: {pdf_folder}")


if __name__ == "__main__":
    main()
