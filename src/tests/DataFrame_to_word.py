from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def save_as_pretty_word(df, output_path):
    # 创建 Word 文档对象
    doc = Document()

    # 设置标题
    title = doc.add_heading("信息论第二章测试复习题库", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 将 Polars DataFrame 转为字典列表方便循环
    records = df.to_dicts()

    for i, row in enumerate(records, 1):
        # 1. 写入题干 (加粗)
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {row['题干']}")
        run.bold = True
        run.font.size = Pt(12)

        # 2. 写入选项 (A, B, C, D 分行)
        for opt in ["A", "B", "C", "D"]:
            col_name = f"选项{opt}"
            if row[col_name]:  # 确保选项不为空
                opt_p = doc.add_paragraph(style="List Bullet")  # 使用列表样式
                opt_p.add_run(f"{opt}. {row[col_name]}")

        # 3. 写入正确答案 (设置颜色区别于题目)
        ans_p = doc.add_paragraph()
        ans_run = ans_p.add_run(f"【正确答案】：{row['正确答案']}")
        ans_run.font.color.rgb = RGBColor(0, 102, 204)  # 深蓝色
        ans_run.font.size = Pt(10)

        # 每个题目之间加一条分割线或空行
        doc.add_paragraph("-" * 112)

    # 保存文档
    doc.save(output_path)
    print(f"Word文件已生成: {output_path}")


# 调用方法
# word_output = r"C:\Users\asus\Desktop\学校作业\kaishi\信息论复习手册.docx"
# save_as_pretty_word(ans, word_output)