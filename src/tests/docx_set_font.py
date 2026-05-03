from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, font_name: str, font_size: int = 12, bold: bool = False):
    """
    设置 run 的字体（同时支持中西文）

    :param run: docx.text.run.Run 对象
    :param font_name: 字体名称（如 "微软雅黑"、"霞鹜文楷"）
    :param font_size: 字号（磅值，12 磅约等于小四）
    :param bold: 是否加粗
    """
    # 1. 设置西文字体
    run.font.name = font_name

    # 2. 设置中文字体（东亚字体）- 安全方式
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.rPr = rPr

    # 获取或创建 rFonts 元素
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    # 设置东亚字体
    rFonts.set(qn("w:eastAsia"), font_name)

    # 3. 设置字号
    run.font.size = Pt(font_size)

    # 4. 设置加粗
    run.font.bold = bold


def process_document(
    input_path: str,
    output_path: str | None = None,
    font_name: str = "微软雅黑",
    font_size: int = 12,
    bold: bool = True,
) -> bool:
    """
    批量处理 Word 文档的字体格式

    :param input_path: 输入文档路径
    :param output_path: 输出文档路径（不指定则覆盖原文件）
    :param font_name: 字体名称
    :param font_size: 字号（磅值）
    :param bold: 是否加粗
    :return: 处理是否成功
    """
    try:
        # 打开文档
        doc = Document(input_path)

        # 遍历所有段落中的所有 run
        for para in doc.paragraphs:
            for run in para.runs:
                set_run_font(run, font_name, font_size, bold)

        # 处理表格中的文本（如果有）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_run_font(run, font_name, font_size, bold)

        # 保存文档
        save_path = output_path if output_path else input_path
        doc.save(save_path)
        print(f"处理完成，已保存至: {save_path}")
        return True

    except Exception as e:
        print(f"处理文档时出错: {e}")
        return False


# 使用示例
if __name__ == "__main__":
    input_file = r"C:\Users\asus\Desktop\学校作业\kaishi\信息论复习手册.docx"
    output_file = r"C:\Users\asus\Desktop\学校作业\kaishi\信息论复习手册_格式化.docx"

    process_document(
        input_path=input_file,
        output_path=output_file,  # 不指定则覆盖原文件
        font_name="霞鹜文楷",  # 字体名称
        font_size=12,  # 字号（12磅 ≈ 小四）
        bold=False,  # 是否加粗（设为 False 保持原样）
    )
