import polars as pl
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_global_font(doc, font_name):
    """设置文档全局中西文字体"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    # 获取或创建底层 XML 节点以支持中文字体
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def save_as_pretty_word(df, output_path, title_text="通信原理填空题题库"):
    """将填空题DataFrame转换为格式美观的Word"""
    doc = Document()
    set_global_font(doc, "霞鹜文楷")

    # 1. 写入大标题
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.name = "霞鹜文楷"
    run.font.size = Pt(18)
    run.font.bold = True
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "霞鹜文楷")

    # 添加副标题/统计信息
    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stats_run = stats_p.add_run(f"共 {len(df)} 道填空题")
    stats_run.font.size = Pt(11)
    stats_run.font.color.rgb = RGBColor(100, 100, 100)

    # 添加分隔线
    doc.add_paragraph("=" * 50)

    records = df.to_dicts()

    for i, row in enumerate(records, 1):
        # 获取题目和答案
        question_text = row.get("题目", "")
        answer_text = row.get("填空题答案", "")

        if not question_text:
            continue

        # A. 题号 + 题干
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {question_text}")
        run.bold = True
        run.font.size = Pt(12)

        # B. 答案（填空题答案）
        ans_p = doc.add_paragraph()
        ans_run = ans_p.add_run(f"【答案】：{answer_text}")
        ans_run.font.color.rgb = RGBColor(192, 0, 0)  # 红色
        ans_run.font.size = Pt(11)
        ans_run.font.bold = True

        # C. 添加一些间距
        doc.add_paragraph()

        # D. 分割线（不是最后一道题才加）
        if i < len(records):
            doc.add_paragraph("-" * 80)

    doc.save(output_path)
    print(f"✅ 已保存: {output_path}")


def batch_process_folder(source_dir, output_dir):
    """
    遍历 source_dir 下所有 Excel，转换并保存到 output_dir
    """
    src_path = Path(source_dir)
    out_path = Path(output_dir)

    # 创建输出文件夹
    out_path.mkdir(parents=True, exist_ok=True)

    # 筛选所有 .xlsx 文件
    files = list(src_path.glob("*.xlsx"))

    if not files:
        print(f"❌ 错误: 在路径 {source_dir} 下没找到 .xlsx 文件")
        return

    print(f"🚀 开始转换任务，共 {len(files)} 个文件...")

    for file in files:
        try:
            # 1. 读取 Excel
            df = pl.read_excel(file)

            # 检查必要的列是否存在
            required_cols = ["题目", "填空题答案"]
            missing_cols = [col for col in required_cols if col not in df.columns]
            if missing_cols:
                print(f"⚠️ 文件 {file.name} 缺少列: {missing_cols}，跳过")
                continue

            # 2. 确定输出路径和文档标题（使用原文件名）
            file_stem = file.stem
            target_word = out_path / f"{file_stem}.docx"

            # 3. 执行转换
            save_as_pretty_word(df, target_word, title_text=file_stem)
            print(f"✅ 已完成: {file_stem}.docx")

        except Exception as e:
            print(f"⚠️ 处理文件 {file.name} 时发生异常: {e}")


def convert_single_file(excel_path, output_path=None, title=None):
    """
    转换单个Excel文件

    参数:
        excel_path: Excel文件路径
        output_path: 输出Word路径（可选，默认同目录同名）
        title: 文档标题（可选，默认使用文件名）
    """
    excel_path = Path(excel_path)

    if not excel_path.exists():
        print(f"❌ 错误: 文件 {excel_path} 不存在")
        return

    # 读取Excel
    df = pl.read_excel(excel_path)

    # 检查必要的列是否存在
    required_cols = ["题目", "填空题答案"]
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        print(f"❌ 错误: Excel缺少列: {missing_cols}")
        return

    # 确定输出路径
    if output_path is None:
        output_dir = excel_path.parent / "word_output"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{excel_path.stem}.docx"
    else:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

    # 确定标题
    if title is None:
        title = excel_path.stem

    # 转换
    save_as_pretty_word(df, output_path, title_text=title)


# ========== 使用示例 ==========
if __name__ == "__main__":
    # 方式1：批量处理整个文件夹
    batch_process_folder(
        source_dir=r"C:\Users\asus\Desktop\学校作业\通信原理",  # Excel文件夹路径
        output_dir=r"C:\Users\asus\Desktop\学校作业\通信原理\word_output",  # 输出Word文件夹路径
    )
