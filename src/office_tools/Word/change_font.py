from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_run_font(run, font_name):
    """设置单个run的字体"""
    # 设置字体名称
    run.font.name = font_name

    # 通过XML设置中文字体和其他语言字体
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')

    # 设置各种语言的字体
    rFonts.set(qn('w:ascii'), font_name)  # 英文字体
    rFonts.set(qn('w:hAnsi'), font_name)  # ANSI字体
    rFonts.set(qn('w:eastAsia'), font_name)  # 东亚字体（中文）
    rFonts.set(qn('w:cs'), font_name)  # 复杂文种字体

    # 替换或添加rFonts元素
    old_rFonts = rPr.find(qn('w:rFonts'))
    if old_rFonts is not None:
        rPr.remove(old_rFonts)
    rPr.append(rFonts)


def convert_word_font(input_path, output_path, font_name="霞鹜文楷"):
    """
    转换Word文档的所有字体为指定字体（修正版）
    """
    print(f"开始处理文档: {input_path}")
    print(f"目标字体: {font_name}")

    # 打开文档
    doc = Document(input_path)

    # 1. 转换所有样式中的字体
    print("正在转换样式字体...")
    for style in doc.styles:
        try:
            if hasattr(style, 'font'):
                style.font.name = font_name

                # 通过XML设置样式字体
                rPr = style._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:cs'), font_name)

                # 替换或添加rFonts元素
                old_rFonts = rPr.find(qn('w:rFonts'))
                if old_rFonts is not None:
                    rPr.remove(old_rFonts)
                rPr.append(rFonts)
        except:
            continue

    # 2. 转换所有段落的字体
    print("正在转换段落字体...")
    para_count = 0
    for paragraph in doc.paragraphs:
        try:
            for run in paragraph.runs:
                set_run_font(run, font_name)
            para_count += 1
            if para_count % 50 == 0:
                print(f"  已处理 {para_count} 个段落...")
        except:
            continue

    # 3. 转换表格中的字体
    print("正在转换表格字体...")
    table_count = 0
    for table in doc.tables:
        try:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, font_name)
            table_count += 1
        except:
            continue

    # 4. 转换页眉页脚中的字体
    print("正在转换页眉页脚字体...")
    try:
        for section in doc.sections:
            # 页眉
            header = section.header
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name)

            # 页脚
            footer = section.footer
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name)
    except Exception as e:
        print(f"  页眉页脚处理提示: {e}")

    # 5. 转换脚注和尾注中的字体（如果存在）
    try:
        if hasattr(doc, 'footnotes'):
            for footnote in doc.footnotes:
                for paragraph in footnote.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, font_name)
    except:
        pass

    # 保存文档
    doc.save(output_path)
    print(f"✓ 文档转换完成！")
    print(f"  处理段落数: {para_count}")
    print(f"  处理表格数: {table_count}")
    print(f"  保存路径: {output_path}")

    return True


def quick_convert_word_font(input_path, output_path, font_name="霞鹜文楷"):
    """
    快速转换Word文档字体（最简版本，只处理基本内容）
    适用于普通文档，速度更快
    """
    print(f"开始处理文档: {input_path}")

    # 打开文档
    doc = Document(input_path)

    # 1. 设置Normal样式
    if "Normal" in doc.styles:
        style = doc.styles["Normal"]
        style.font.name = font_name
        try:
            rPr = style._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            old_rFonts = rPr.find(qn('w:rFonts'))
            if old_rFonts is not None:
                rPr.remove(old_rFonts)
            rPr.append(rFonts)
        except:
            pass

    # 2. 转换所有段落的run
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            try:
                run.font.name = font_name
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:ascii'), font_name)
                old_rFonts = rPr.find(qn('w:rFonts'))
                if old_rFonts is not None:
                    rPr.remove(old_rFonts)
                rPr.append(rFonts)
            except:
                continue

    # 3. 转换表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        try:
                            run.font.name = font_name
                            rPr = run._element.get_or_add_rPr()
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:eastAsia'), font_name)
                            rFonts.set(qn('w:ascii'), font_name)
                            old_rFonts = rPr.find(qn('w:rFonts'))
                            if old_rFonts is not None:
                                rPr.remove(old_rFonts)
                            rPr.append(rFonts)
                        except:
                            continue

    # 保存文档
    doc.save(output_path)
    print(f"✓ 快速转换完成！保存至: {output_path}")

    return True


# 使用示例
if __name__ == "__main__":
    # 完整版转换
    # convert_word_font(
    #     input_path=r"F:\E\编程学习\01 编程语言学习\Rust\Rust 程序设计语言 简体中文版.docx",
    #     output_path=r"F:\E\编程学习\01 编程语言学习\Rust\Rust 程序设计语言_新字体.docx",
    #     font_name="霞鹜文楷等宽 Medium"
    # )

    # 如果完整版还有问题，可以使用快速版
    quick_convert_word_font(
        input_path=r"F:\E\编程学习\01 编程语言学习\Rust\Rust 程序设计语言 简体中文版.docx",
        output_path=r"F:\E\编程学习\01 编程语言学习\Rust\Rust 程序设计语言_新字体.docx",
        font_name="霞鹜文楷等宽 Medium"
    )