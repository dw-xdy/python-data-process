import polars as pl
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# --- 代码是将 excel 中的数据批量导入 word 并整理好格式的代码 ---
def set_global_font(doc, font_name):
    """设置文档全局中西文字体"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    # 获取或创建底层 XML 节点以支持中文字体
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def save_as_pretty_word(df, output_path, title_text="复习题库"):
    """将单个 DataFrame 转换为格式美观的 Word"""
    doc = Document()
    set_global_font(doc, "霞鹜文楷")

    # 1. 写入大标题
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.name = "霞鹜文楷"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "霞鹜文楷")

    records = df.to_dicts()

    for i, row in enumerate(records, 1):
        # A. 题干 (加粗, 12pt)
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {row['题干']}")
        run.bold = True
        run.font.size = Pt(12)

        # B. 选项 (A,B,C,D)
        for opt in ["A", "B", "C", "D"]:
            col_name = f"选项{opt}"
            # 检查列是否存在且不为空
            if col_name in row and row[col_name]:
                opt_p = doc.add_paragraph(style="List Bullet")
                opt_run = opt_p.add_run(f"{opt}. {row[col_name]}")
                opt_run.font.bold = True

        # C. 正确答案 (深蓝色, 10pt)
        ans_p = doc.add_paragraph()
        ans_run = ans_p.add_run(f"【正确答案】：{row['正确答案']}")
        ans_run.font.color.rgb = RGBColor(0, 102, 204)
        ans_run.font.size = Pt(10)
        ans_run.font.bold = True

        # D. 分割线
        doc.add_paragraph("-" * 80)

    doc.save(output_path)


# --- 批量处理逻辑 ---


def batch_process_folder(source_dir, output_dir):
    """
    遍历 source_dir 下所有 Excel，转换并保存到 output_dir
    """
    src_path = Path(source_dir)
    out_path = Path(output_dir)

    # 创建输出文件夹
    out_path.mkdir(parents=True, exist_ok=True)

    # 筛选所有 .xlsx 文件
    files = list(src_path.glob("*.xlsx"))

    if not files:
        print(f"❌ 错误: 在路径 {source_dir} 下没找到 .xlsx 文件")
        return

    print(f"🚀 开始转换任务，共 {len(files)} 个文件...")

    for file in files:
        try:
            # 1. 读取 Excel
            df = pl.read_excel(file)

            # 2. 确定输出路径和文档标题
            file_stem = file.stem
            target_word = out_path / f"{file_stem}.docx"

            # 3. 执行转换
            save_as_pretty_word(df, target_word, title_text=file_stem)
            print(f"✅ 已完成: {file_stem}.docx")

        except Exception as e:
            print(f"⚠️ 处理文件 {file.name} 时发生异常: {e}")
