import polars as pl
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_global_font(doc, font_name):
    """设置文档全局中西文字体"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)

    # 获取或创建底层 XML 节点以支持中文字体
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def save_as_pretty_word(df, output_path, title_text="复习题库"):
    """将单个 DataFrame 转换为格式美观的 Word"""
    doc = Document()
    set_global_font(doc, "霞鹜文楷")

    # 1. 写入大标题
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.font.name = "霞鹜文楷"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "霞鹜文楷")

    records = df.to_dicts()

    for i, row in enumerate(records, 1):
        # A. 题干（题目列）
        # 注意：列名是"题目"，不是"题干"
        question_text = row.get("题目", "")
        if not question_text:
            continue

        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {question_text}")
        run.bold = True
        run.font.size = Pt(12)

        # B. 选项（A,B,C,D）- 只显示非空的选项
        for opt in ["A", "B", "C", "D"]:
            col_name = f"选项{opt}"
            opt_value = row.get(col_name, "")
            if opt_value and str(opt_value).strip():
                opt_p = doc.add_paragraph(style="List Bullet")
                opt_run = opt_p.add_run(f"{opt}. {opt_value}")
                opt_run.font.bold = True

        # C. 正确答案
        # 注意：答案列是"答案"，值为1表示正确，0表示错误
        raw_answer = row.get("答案", "")
        # 将数字答案转换为可读文本
        if raw_answer == 1 or str(raw_answer).strip() == "1":
            answer_text = "正确"
        elif raw_answer == 0 or str(raw_answer).strip() == "0":
            answer_text = "错误"
        else:
            answer_text = str(raw_answer) if raw_answer else ""

        ans_p = doc.add_paragraph()
        ans_run = ans_p.add_run(f"【正确答案】：{answer_text}")
        ans_run.font.color.rgb = RGBColor(0, 102, 204)
        ans_run.font.size = Pt(10)
        ans_run.font.bold = True

        # D. 分割线
        doc.add_paragraph("-" * 80)

    doc.save(output_path)
    print(f"✅ 已保存: {output_path}")


def batch_process_folder(source_dir, output_dir):
    """
    遍历 source_dir 下所有 Excel，转换并保存到 output_dir
    """
    src_path = Path(source_dir)
    out_path = Path(output_dir)

    # 创建输出文件夹
    out_path.mkdir(parents=True, exist_ok=True)

    # 筛选所有 .xlsx 文件
    files = list(src_path.glob("*.xlsx"))

    if not files:
        print(f"❌ 错误: 在路径 {source_dir} 下没找到 .xlsx 文件")
        return

    print(f"🚀 开始转换任务，共 {len(files)} 个文件...")

    for file in files:
        try:
            # 1. 读取 Excel
            df = pl.read_excel(file)

            # 2. 确定输出路径和文档标题（使用原文件名）
            file_stem = file.stem
            target_word = out_path / f"{file_stem}.docx"

            # 3. 执行转换
            save_as_pretty_word(df, target_word, title_text=file_stem)
            print(f"✅ 已完成: {file_stem}.docx")

        except Exception as e:
            print(f"⚠️ 处理文件 {file.name} 时发生异常: {e}")


# ========== 使用示例 ==========
if __name__ == "__main__":
    # 方式1：批量处理整个文件夹
    batch_process_folder(
        source_dir=r"C:\Users\asus\Desktop\学校作业\通信原理",  # 替换为你的Excel文件夹路径
        output_dir=r"C:\Users\asus\Desktop\学校作业\通信原理\an_word",  # 替换为输出Word文件夹路径
    )
